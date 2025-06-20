# ==== Free Plan ====



import openai
import PyPDF2
import docx
from google.colab import files


openai.api_key = st.secrets["OPENAI_API_KEY"]


def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text



def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])



def analyze_resume_basic(resume_text, job_description=""):
    prompt = (
        "You are a resume expert. Analyze the resume below and provide:\n"
        "- ATS compatibility score out of 100\n"
        "- Matching and missing keywords (based on job description)\n"
        "- General feedback (clarity, relevance, structure)\n\n"
        f"RESUME:\n{resume_text}\n\n"
        f"JOB DESCRIPTION:\n{job_description}"
    )
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()



uploaded = files.upload()
file_name = list(uploaded.keys())[0]

if file_name.endswith(".pdf"):
    resume_text = extract_text_from_pdf(file_name)
elif file_name.endswith(".docx"):
    resume_text = extract_text_from_docx(file_name)
else:
    raise ValueError("Unsupported file format. Upload .pdf or .docx")


job_description = input("Paste the job description here (optional):\n")


print("🔍 Analyzing your resume...")
feedback = analyze_resume_basic(resume_text, job_description)
print("\n📝 Feedback:\n")
print(feedback)


# ==== Pro Plan ====



import openai
import PyPDF2
import docx
from fpdf import FPDF
from google.colab import files


openai.api_key = st.secrets["OPENAI_API_KEY"]


def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])



def analyze_resume(resume_text, job_description=""):
    prompt = (
        "You are a resume expert. Analyze the resume below and provide:\n"
        "- ATS compatibility score out of 100\n"
        "- Matching and missing keywords (based on job description)\n"
        "- General feedback (clarity, relevance, structure)\n\n"
        f"RESUME:\n{resume_text}\n\n"
        f"JOB DESCRIPTION:\n{job_description}"
    )
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()



def rewrite_resume(resume_text):
    prompt = (
        "You're an expert resume writer. Rewrite the resume below to be more professional, concise, and impactful, while preserving the original meaning.\n\n"
        f"RESUME:\n{resume_text}"
    )
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()



def export_to_pdf(text, filename="Rewritten_Resume.pdf"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("helvetica", size=11)
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf.output(filename)
    files.download(filename)



uploaded = files.upload()
file_name = list(uploaded.keys())[0]

if file_name.endswith(".pdf"):
    resume_text = extract_text_from_pdf(file_name)
elif file_name.endswith(".docx"):
    resume_text = extract_text_from_docx(file_name)
else:
    raise ValueError("Unsupported file format. Upload .pdf or .docx")


job_description = input("Paste the job description here (optional):\n")


print("🔍 Analyzing resume...")
feedback = analyze_resume(resume_text, job_description)
print("\n📝 Analysis Result:\n")
print(feedback)



rewrite_confirm = input("Do you want a full rewritten version of this resume? (yes/no): ").strip().lower()
if rewrite_confirm == "yes":
    rewritten_resume = rewrite_resume(resume_text)
    print("\n✅ Rewritten Resume:\n")
    print(rewritten_resume)
    save_pdf = input("Export rewritten resume as PDF? (yes/no): ").strip().lower()
    if save_pdf == "yes":
        export_to_pdf(rewritten_resume)


# ==== Unlimited Plan ====



import openai
import PyPDF2
import docx
from fpdf import FPDF
from google.colab import files


openai.api_key = st.secrets["OPENAI_API_KEY"]


def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])



def analyze_resume(resume_text, job_description=""):
    prompt = (
        "You are a resume expert. Analyze the resume below and provide:\n"
        "- ATS compatibility score out of 100\n"
        "- Matching and missing keywords (based on job description)\n"
        "- General feedback (clarity, relevance, structure)\n\n"
        f"RESUME:\n{resume_text}\n\n"
        f"JOB DESCRIPTION:\n{job_description}"
    )
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()



def rewrite_resume(resume_text):
    prompt = (
        "You're an expert resume writer. Rewrite the resume below to be more professional, concise, and impactful, while preserving the original meaning.\n\n"
        f"RESUME:\n{resume_text}"
    )
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()



def export_to_pdf(text, filename="Rewritten_Resume.pdf"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("helvetica", size=11)
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf.output(filename)
    files.download(filename)



while True:
    print("\n--- NEW RESUME SCAN ---")
    uploaded = files.upload()

    if not uploaded:
        print("❌ No file uploaded. Please try again.")
        continue

    file_name = list(uploaded.keys())[0]

    if file_name.endswith(".pdf"):
        resume_text = extract_text_from_pdf(file_name)
    elif file_name.endswith(".docx"):
        resume_text = extract_text_from_docx(file_name)
    else:
        print("❌ Unsupported file format. Use .pdf or .docx only.")
        continue

    job_description = input("Paste the job description here (optional):\n")

    print("🔍 Analyzing resume...")
    feedback = analyze_resume(resume_text, job_description)
    print("\n📝 Analysis Result:\n")
    print(feedback)

    rewrite_confirm = input("Do you want a full rewritten version of this resume? (yes/no): ").strip().lower()
    if rewrite_confirm == "yes":
        rewritten_resume = rewrite_resume(resume_text)
        print("\n✅ Rewritten Resume:\n")
        print(rewritten_resume)

        save_pdf = input("Export rewritten resume as PDF? (yes/no): ").strip().lower()
        if save_pdf == "yes":
            export_to_pdf(rewritten_resume)

    another = input("Do you want to scan another resume? (yes/no): ").strip().lower()
    if another != "yes":
        print("✅ Session ended. Thank you for using the Unlimited Plan.")
        break


